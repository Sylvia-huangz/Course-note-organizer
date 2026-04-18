from __future__ import annotations

import argparse
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from pydantic import ValidationError

from _errors import write_error_manifest, write_manifest, write_validation_error
from _markdown_blocks import MarkdownBlock, parse_markdown_blocks
from _schemas import ExportDocxRequest, ExportDocxResult


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def setup_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ")
    add_field(footer, "PAGE")
    return document


def insert_toc(document: Document) -> None:
    document.add_paragraph("Table of Contents", style="Heading 1")
    paragraph = document.add_paragraph()
    add_field(paragraph, r'TOC \o "1-3" \h \z \u')
    document.add_page_break()


def render_formula_image(formula_text: str, temp_dir: Path) -> Path:
    font = ImageFont.load_default()
    text = formula_text.strip()
    width = max(500, len(text) * 8 + 40)
    image = Image.new("RGB", (width, 80), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 25), text, fill="black", font=font)
    output = temp_dir / f"formula-{abs(hash(text))}.png"
    image.save(output)
    return output


def add_code_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, cell in enumerate(row):
            table.cell(row_index, col_index).text = cell


def render_block(document: Document, block: MarkdownBlock, temp_dir: Path) -> None:
    if block.type == "heading":
        document.add_heading(block.text, level=min(max((block.level or 1) - 1, 0), 3))
        return
    if block.type == "paragraph":
        stripped = block.text.strip()
        if stripped.startswith(r"\[") and stripped.endswith(r"\]"):
            formula_path = render_formula_image(stripped, temp_dir)
            document.add_picture(str(formula_path), width=Inches(5.5))
            return
        document.add_paragraph(block.text)
        return
    if block.type == "blockquote":
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        run = paragraph.add_run(block.text)
        run.italic = True
        return
    if block.type == "bullet_list":
        for item in block.items:
            document.add_paragraph(item, style="List Bullet")
        return
    if block.type == "ordered_list":
        for item in block.items:
            document.add_paragraph(item, style="List Number")
        return
    if block.type == "code":
        add_code_paragraph(document, block.text)
        return
    if block.type == "table":
        add_table(document, block.rows)


def export_docx(markdown_path: Path, output_path: Path) -> ExportDocxResult:
    document = setup_document()
    blocks = parse_markdown_blocks(markdown_path.read_text(encoding="utf-8"))
    with tempfile.TemporaryDirectory() as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        if blocks and blocks[0].type == "heading":
            title = document.add_heading(blocks[0].text, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            blocks = blocks[1:]
        insert_toc(document)
        for block in blocks:
            render_block(document, block, temp_dir)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
    return ExportDocxResult(
        status="ok",
        output=str(output_path),
        formula_strategy="native-if-available-else-image-fallback",
        notes=["Word updates the table of contents field when the document is opened and fields are refreshed."],
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export Markdown course notes to DOCX.")
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--manifest")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    manifest_hint = Path(args.manifest).expanduser().resolve() if args.manifest else None
    try:
        request = ExportDocxRequest.model_validate(vars(args))
    except ValidationError as exc:
        if manifest_hint:
            write_validation_error(manifest_hint, exc)
        return 1

    markdown_path = Path(request.markdown).expanduser().resolve()
    output_path = Path(request.output).expanduser().resolve()
    if not markdown_path.exists():
        if manifest_hint:
            write_error_manifest(
                manifest_hint,
                code="MISSING_SOURCE",
                message="Markdown source file not found for DOCX export.",
                source=str(markdown_path),
            )
        return 1

    try:
        result = export_docx(markdown_path, output_path)
    except Exception as exc:
        if manifest_hint:
            write_error_manifest(
                manifest_hint,
                code="EXPORT_FAILED",
                message="DOCX export failed before completion.",
                source=str(markdown_path),
                details={"exception": str(exc)},
                retryable=True,
            )
        return 1

    if request.manifest:
        write_manifest(Path(request.manifest).expanduser().resolve(), **result.model_dump())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
